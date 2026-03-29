#!/usr/bin/env python3
"""
Report Generator

Generates professional DOCX market research reports with charts,
tables, and comprehensive analysis.

Kimi-Specific Implementation:
- Uses python-docx for document generation
- Integrates matplotlib for chart creation
- Demonstrates structured document output
"""

from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import matplotlib.pyplot as plt
import os

from .chart_generator import ChartGenerator


class ReportGenerator:
    """
    Generates professional market research reports.
    
    Report Sections:
    - Executive Summary
    - Competitor Analysis
    - Pricing Intelligence
    - Sentiment Analysis
    - Market Trends
    - Strategic Recommendations
    """
    
    def __init__(self):
        self.chart_generator = ChartGenerator()
        self.temp_charts = []
    
    def generate(self, 
                 results: Any, 
                 output_path: str = "market_research_report.docx") -> str:
        """
        Generate complete market research report.
        
        Args:
            results: ResearchResults object with all data
            output_path: Path for output DOCX file
            
        Returns:
            Path to generated report
        """
        print("   📝 Creating report sections...")
        
        # Create document
        doc = Document()
        
        # Add title page
        self._add_title_page(doc, results)
        
        # Add executive summary
        self._add_executive_summary(doc, results)
        
        # Add competitor analysis
        self._add_competitor_analysis(doc, results.competitor_analysis)
        
        # Add pricing intelligence
        self._add_pricing_analysis(doc, results.pricing_data)
        
        # Add sentiment analysis
        self._add_sentiment_analysis(doc, results.sentiment_analysis)
        
        # Add market trends
        self._add_market_trends(doc, results.market_trends)
        
        # Add recommendations
        self._add_recommendations(doc, results)
        
        # Save document
        doc.save(output_path)
        
        # Clean up temp charts
        self._cleanup_charts()
        
        return output_path
    
    def _add_title_page(self, doc: Document, results: Any):
        """Add title page to document"""
        # Title
        title = doc.add_heading('Market Research Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"{results.target_company}")
        run.font.size = Pt(18)
        run.font.bold = True
        
        # Industry
        industry = doc.add_paragraph()
        industry.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = industry.add_run(f"{results.industry} Industry")
        run.font.size = Pt(14)
        
        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
        run.font.size = Pt(12)
        
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, results: Any):
        """Add executive summary section"""
        doc.add_heading('Executive Summary', 1)
        
        # Overview
        doc.add_paragraph(
            f"This report provides comprehensive market research analysis for "
            f"{results.target_company} in the {results.industry} industry. "
            f"The analysis covers competitor positioning, pricing intelligence, "
            f"customer sentiment, and market trends."
        )
        
        # Key findings
        doc.add_heading('Key Findings', 2)
        
        findings = [
            f"Analyzed {len(results.competitor_analysis.get('competitors', []))} key competitors",
            f"Reviewed pricing for {results.pricing_data.get('price_analysis', {}).get('total_products', 0)} products",
            f"Analyzed {results.sentiment_analysis.get('review_count', 0):,} customer reviews",
            f"Overall market sentiment: {results.sentiment_analysis.get('overall_sentiment', 'N/A')}"
        ]
        
        for finding in findings:
            doc.add_paragraph(finding, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_competitor_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add competitor analysis section"""
        doc.add_heading('Competitor Analysis', 1)
        
        # Competitor overview table
        doc.add_heading('Competitor Overview', 2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Company'
        hdr_cells[1].text = 'Market Share'
        hdr_cells[2].text = 'Revenue'
        hdr_cells[3].text = 'Employees'
        
        # Data rows
        for comp in data.get('competitors', []):
            row_cells = table.add_row().cells
            row_cells[0].text = comp.get('name', '')
            row_cells[1].text = f"{comp.get('market_share', 0)}%"
            row_cells[2].text = comp.get('revenue', '')
            row_cells[3].text = comp.get('employees', '')
        
        # Market positioning
        doc.add_heading('Market Positioning', 2)
        positioning = data.get('market_positioning', {})
        
        for segment, companies in positioning.items():
            if companies:
                p = doc.add_paragraph()
                p.add_run(f"{segment.replace('_', ' ').title()}: ").bold = True
                p.add_run(', '.join(companies))
        
        # Strategic insights
        doc.add_heading('Strategic Insights', 2)
        for insight in data.get('strategic_insights', []):
            doc.add_paragraph(insight, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_pricing_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add pricing intelligence section"""
        doc.add_heading('Pricing Intelligence', 1)
        
        # Price statistics
        stats = data.get('price_analysis', {}).get('price_statistics', {})
        
        doc.add_heading('Price Overview', 2)
        doc.add_paragraph(f"Average Price: ${stats.get('mean', 0):,.2f}")
        doc.add_paragraph(f"Median Price: ${stats.get('median', 0):,.2f}")
        doc.add_paragraph(f"Price Range: ${stats.get('min', 0):,.0f} - ${stats.get('max', 0):,.0f}")
        
        # Price comparison table
        doc.add_heading('Price Comparison by Category', 2)
        
        comparisons = data.get('price_comparison', {})
        for category, products in comparisons.items():
            doc.add_heading(category.title(), 3)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light List Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Company'
            hdr_cells[1].text = 'Product'
            hdr_cells[2].text = 'Price'
            
            for company, product, price in products:
                row_cells = table.add_row().cells
                row_cells[0].text = company
                row_cells[1].text = product
                row_cells[2].text = f"${price:,.0f}"
        
        # Recommendations
        doc.add_heading('Pricing Recommendations', 2)
        for rec in data.get('recommendations', []):
            doc.add_paragraph(rec, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_sentiment_analysis(self, doc: Document, data: Dict[str, Any]):
        """Add sentiment analysis section"""
        doc.add_heading('Sentiment Analysis', 1)
        
        # Overall sentiment
        doc.add_heading('Overall Market Sentiment', 2)
        doc.add_paragraph(f"Market Sentiment: {data.get('overall_sentiment', 'N/A')}")
        doc.add_paragraph(f"Total Reviews Analyzed: {data.get('review_count', 0):,}")
        
        # Company sentiment comparison
        doc.add_heading('Sentiment by Company', 2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Company'
        hdr_cells[1].text = 'Positive'
        hdr_cells[2].text = 'Neutral'
        hdr_cells[3].text = 'Negative'
        
        sentiments = data.get('company_sentiments', {})
        for company, sentiment in sentiments.items():
            row_cells = table.add_row().cells
            row_cells[0].text = company
            row_cells[1].text = f"{sentiment.get('positive', 0)}%"
            row_cells[2].text = f"{sentiment.get('neutral', 0)}%"
            row_cells[3].text = f"{sentiment.get('negative', 0)}%"
        
        # Key themes
        doc.add_heading('Key Themes', 2)
        themes = data.get('key_themes', {})
        
        doc.add_paragraph('Top Positive Themes:').bold = True
        for theme in themes.get('top_positive', [])[:5]:
            doc.add_paragraph(theme, style='List Bullet')
        
        doc.add_paragraph('Top Negative Themes:').bold = True
        for theme in themes.get('top_negative', [])[:5]:
            doc.add_paragraph(theme, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_market_trends(self, doc: Document, data: Dict[str, Any]):
        """Add market trends section"""
        doc.add_heading('Market Trends', 1)
        
        # Industry growth
        doc.add_heading('Industry Growth', 2)
        doc.add_paragraph(f"Industry Growth Rate: {data.get('industry_growth', 'N/A')}")
        doc.add_paragraph(f"Market Size: {data.get('market_size', 'N/A')}")
        doc.add_paragraph(f"Growth Forecast: {data.get('growth_forecast', 'N/A')}")
        
        # Key trends
        doc.add_heading('Key Market Trends', 2)
        for trend in data.get('key_trends', []):
            doc.add_paragraph(trend, style='List Bullet')
        
        doc.add_page_break()
    
    def _add_recommendations(self, doc: Document, results: Any):
        """Add strategic recommendations section"""
        doc.add_heading('Strategic Recommendations', 1)
        
        recommendations = [
            {
                "title": "Competitive Positioning",
                "content": f"Based on the analysis, {results.target_company} should focus on "
                          f"differentiating through innovation and customer experience to "
                          f"maintain competitive advantage."
            },
            {
                "title": "Pricing Strategy",
                "content": "Consider value-based pricing that reflects the unique features "
                          "and benefits offered. Monitor competitor pricing closely for "
                          "market adjustments."
            },
            {
                "title": "Customer Engagement",
                "content": "Leverage positive sentiment themes in marketing. Address "
                          "negative feedback proactively to improve customer satisfaction."
            },
            {
                "title": "Market Opportunities",
                "content": "The market shows strong growth potential. Consider expansion "
                          "into underserved segments and geographic markets."
            }
        ]
        
        for rec in recommendations:
            doc.add_heading(rec['title'], 2)
            doc.add_paragraph(rec['content'])
        
        # Conclusion
        doc.add_heading('Conclusion', 1)
        doc.add_paragraph(
            f"The {results.industry} market presents significant opportunities for "
            f"{results.target_company}. By leveraging strengths in innovation and "
            f"addressing areas for improvement, the company can strengthen its "
            f"market position and drive growth."
        )
    
    def _cleanup_charts(self):
        """Remove temporary chart files"""
        for chart_path in self.temp_charts:
            try:
                if os.path.exists(chart_path):
                    os.remove(chart_path)
            except:
                pass


# Example usage
if __name__ == "__main__":
    from dataclasses import dataclass, field
    
    @dataclass
    class MockResults:
        target_company: str = "Tesla"
        industry: str = "Electric Vehicles"
        competitor_analysis: Dict = field(default_factory=dict)
        pricing_data: Dict = field(default_factory=dict)
        sentiment_analysis: Dict = field(default_factory=dict)
        market_trends: Dict = field(default_factory=dict)
    
    generator = ReportGenerator()
    
    # Create mock results
    results = MockResults()
    
    # Generate report
    path = generator.generate(results, "test_report.docx")
    print(f"Report generated: {path}")
